# app.py
"""
KTPS 10-day Report Dashboard (Streamlit)
- Upload KTPS 10-day .docx reports
- Auto-parse key fields (VC, specific coal cons, coal sources, reasons breakup)
- Let user review/edit parsed values before saving
- Append record to local CSV (data/ktps_10day_reports.csv)
- Optionally append to Google Sheets (set USE_GSHEETS = True and configure Streamlit secrets)
"""

import streamlit as st
import pandas as pd
import io, re, os, json
from datetime import datetime
from docx import Document
import plotly.express as px

# ---------- CONFIG ----------
# If True, app will attempt to append records to Google Sheets using service account info stored in
# Streamlit secrets: st.secrets["gcp_service_account"] (JSON/dict) and st.secrets["gspread_sheet_id"] (sheet id)
USE_GSHEETS = True

DATA_DIR = "data"
os.makedirs(DATA_DIR, exist_ok=True)
DATA_FILE = os.path.join(DATA_DIR, "ktps_10day_reports.csv")

# Try to import gspread/google auth if available
try:
    import gspread
    from google.oauth2.service_account import Credentials
    GSPREAD_AVAILABLE = True
except Exception:
    GSPREAD_AVAILABLE = False

st.set_page_config(page_title="KTPS 10-day Dashboard", layout="wide")
st.title("KTPS — 10-Day Report Dashboard (Parser + Append + Monthly Report)")

# ---------- Parsing helpers ----------
def extract_text_from_docx_bytes(b: bytes) -> str:
    """Extract text (paragraphs + table rows) from a .docx file bytes."""
    doc = Document(io.BytesIO(b))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)

def find_first_float(patterns, text):
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            num = re.search(r"-?\d+[\d,]*\.?\d*", m.group())
            if num:
                return float(num.group().replace(",", ""))
    return None

def parse_variable_charge(text: str) -> dict:
    lines = text.splitlines()
    vc = {"merc_order": None, "recoverable_mod": None, "actual": None,
          "disallowance_rs_cr": None, "unrecoverable_rpkwh": None}
    for line in lines:
        low = line.lower()
        if "variable charge" in low and "merc" in low:
            m = re.search(r"(-?\d+[\d,]*\.?\d*)", line)
            if m: vc["merc_order"] = float(m.group().replace(",", ""))
        if "variable charge (recoverable" in low or ("recoverable" in low and "mod" in low):
            m = re.search(r"(-?\d+[\d,]*\.?\d*)", line)
            if m: vc["recoverable_mod"] = float(m.group().replace(",", ""))
        if "variable charge (actual" in low or "actual energy charge" in low:
            m = re.search(r"(-?\d+[\d,]*\.?\d*)", line)
            if m: vc["actual"] = float(m.group().replace(",", ""))
        if "disallowance" in low and "rs" in low:
            m = re.search(r"-?\d+[\d,]*\.?\d*", line)
            if m: vc["disallowance_rs_cr"] = float(m.group().replace(",", ""))
        if "unrecoverable" in low:
            m = re.search(r"-?\d+[\d,]*\.?\d*", line)
            if m: vc["unrecoverable_rpkwh"] = float(m.group().replace(",", ""))
    # fallback regex searches
    if vc["merc_order"] is None:
        vc["merc_order"] = find_first_float([r"variable charge.*merc.*rs\/?kwh[:\s]*([0-9\.\, -]+)"], text)
    if vc["recoverable_mod"] is None:
        vc["recoverable_mod"] = find_first_float([r"variable charge.*recoverable.*rs\/?kwh[:\s]*([0-9\.\, -]+)"], text)
    if vc["actual"] is None:
        vc["actual"] = find_first_float([r"variable charge.*actual.*rs\/?kwh[:\s]*([0-9\.\, -]+)", r"actual energy charge.*rs\/?kwh[:\s]*([0-9\.\, -]+)"], text)
    if vc["disallowance_rs_cr"] is None:
        m = re.search(r"disallowance\s*rs[\. ]*cr[:\s]*(-?\d+[\d,]*\.?\d*)", text, re.IGNORECASE)
        if m:
            vc["disallowance_rs_cr"] = float(m.group(1).replace(",", ""))
    return vc

def parse_sp_coal(text: str) -> dict:
    m = re.search(r"sp coal consumption.*allowable.*actual.*?([\d\.]+).*?([\d\.]+)", text, re.IGNORECASE|re.DOTALL)
    if m:
        return {"allowable": float(m.group(1)), "actual": float(m.group(2))}
    allow = find_first_float([r"allowable.*norm[:\s]*([0-9\.]+)", r"allowable as per merc norm[:\s]*([0-9\.]+)"], text)
    actual = find_first_float([r"actual[:\s]*([0-9]\.[0-9]+)\s*$", r"actual[:\s]*([0-9]\.[0-9]+)"], text)
    return {"allowable": allow, "actual": actual}

def parse_coal_sources(text: str) -> list:
    lines = text.splitlines()
    sources = []
    start = None
    for idx,l in enumerate(lines):
        if re.search(r"coal source", l, re.IGNORECASE) and ("gcv" in l.lower() or "landed cost" in l.lower()):
            start = idx+1
            break
    if start is not None:
        for l in lines[start:start+200]:
            if not l.strip(): break
            nums = re.findall(r"-?\d+[\d,]*\.?\d*", l)
            name = re.sub(r"[-\d\.,\|\t ]+", " ", l).strip()
            try:
                gcv = float(nums[0].replace(",", "")) if len(nums)>=1 else None
                landed = float(nums[1].replace(",", "")) if len(nums)>=2 else None
                eff = float(nums[2].replace(",", "")) if len(nums)>=3 else None
                pct = float(nums[3].replace(",", "")) if len(nums)>=4 else None
            except:
                gcv=landed=eff=pct=None
            sources.append({"name": name, "GCV_ARB": gcv, "Landed_RsPerMT": landed, "Eff_RsPerMkcal": eff, "Pct_share": pct})
    return sources

def parse_reasons_breakup(text: str) -> list:
    reasons = []
    for m in re.finditer(r"(?P<title>[\w \-/\(\)\&]+?)\s*(?:Loss|loss)\s*[-–]\s*(?P<amt>-?\d+[\d,]*\.?\d*)\s*rs\.?cr", text, re.IGNORECASE):
        title = m.group("title").strip()
        amt = float(m.group("amt").replace(",", ""))
        reasons.append({"reason": title, "amount_rs_cr": amt})
    fallback_reasons = ["Diff in LEQ & ARB GCV", "Stack loss", "Plant / O&M /Other losses", "Sp Oil Consn", "Aux power consn"]
    for fr in fallback_reasons:
        if fr.lower() in text.lower() and not any(fr.lower() in r["reason"].lower() for r in reasons):
            m = re.search(fr + r".*?(-?\d+[\d,]*\.?\d*)", text, re.IGNORECASE|re.DOTALL)
            amt = float(m.group(1).replace(",", "")) if m else None
            reasons.append({"reason": fr, "amount_rs_cr": amt})
    return reasons

# ---------- Google Sheets helper ----------
def gsheet_append_row(record: dict):
    """
    Append record to Google Sheet using service account stored in st.secrets:
    - st.secrets["gcp_service_account"] : service account JSON (as dict)
    - st.secrets["gspread_sheet_id"] : sheet id
    Returns (ok:bool, message:str)
    """
    if not GSPREAD_AVAILABLE:
        return False, "gspread/google-auth not installed"
    try:
        secrets = st.secrets
        if "gcp_service_account" not in secrets or "gspread_sheet_id" not in secrets:
            return False, "gcp_service_account or gspread_sheet_id missing in Streamlit secrets"
        sa_info = secrets["gcp_service_account"]
        sheet_id = secrets["gspread_sheet_id"]
        scopes = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(sheet_id)
        ws = sh.sheet1
        header = ["station", "period_label", "report_date", "vc_merc_order", "vc_recoverable_mod",
                  "vc_actual", "vc_unrecoverable_rpkwh", "fuel_cost_disallowance_rs_cr",
                  "sp_coal_allowable", "sp_coal_actual", "aux_power_pct", "sp_oil_ml_per_kwh",
                  "coal_sources", "reasons_breakup", "raw_text_excerpt"]
        try:
            existing_header = ws.row_values(1)
        except Exception:
            existing_header = []
        if not existing_header or len(existing_header) < 3:
            ws.insert_row(header, index=1)
        row = [str(record.get(c, "")) for c in header]
        ws.append_row(row, value_input_option='USER_ENTERED')
        return True, "Appended to Google Sheet"
    except Exception as e:
        return False, str(e)

# ---------- UI: Upload, parse, edit, save ----------
st.header("1) Upload a 10-day KTPS .docx report")
uploaded = st.file_uploader("Upload KTPS .docx (one 10-day 10-day report)", type=["docx"])

if uploaded is not None:
    raw_bytes = uploaded.getvalue()
    raw_text = extract_text_from_docx_bytes(raw_bytes)
    st.subheader("Preview of extracted text (first 900 chars)")
    st.code(raw_text[:900] + ("\n\n... (truncated)" if len(raw_text) > 900 else ""), language="text")

    st.subheader("Auto-parsed fields — review & edit before saving")
    vc = parse_variable_charge(raw_text)
    spcoal = parse_sp_coal(raw_text)
    coal_sources = parse_coal_sources(raw_text)
    reasons = parse_reasons_breakup(raw_text)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Variable Charge (Rs/kWh)**")
        mer_order = st.number_input("MERC Order (Rs/kWh)", value=float(vc.get("merc_order") or 0.0), format="%.4f")
        recoverable = st.number_input("Recoverable (MOD) (Rs/kWh)", value=float(vc.get("recoverable_mod") or 0.0), format="%.4f")
        actual_vc = st.number_input("Actual VC (Rs/kWh)", value=float(vc.get("actual") or 0.0), format="%.4f")
        unrecoverable_rpkwh = st.number_input("Unrecoverable (Rs/kWh)", value=float(vc.get("unrecoverable_rpkwh") or 0.0), format="%.4f")
        disallowance_cr = st.number_input("Fuel cost disallowance (Rs.Cr)", value=float(vc.get("disallowance_rs_cr") or 0.0), format="%.2f")
    with col2:
        st.markdown("**Specific coal consumption (kg/kWh)**")
        sp_allow = st.number_input("Allowable (kg/kWh)", value=float(spcoal.get("allowable") or 0.0), format="%.4f")
        sp_actual = st.number_input("Actual (kg/kWh)", value=float(spcoal.get("actual") or 0.0), format="%.4f")
        st.markdown("**Other station metrics (optional)**")
        aux_power = st.number_input("Aux power (%)", value=0.0, format="%.3f")
        sp_oil = st.number_input("Sp Oil Consn (ml/kWh)", value=0.0, format="%.3f")

    st.markdown("**Coal sources detected** (edit/add rows if parser missed anything)")
    cs_df = pd.DataFrame(coal_sources) if coal_sources else pd.DataFrame(columns=["name","GCV_ARB","Landed_RsPerMT","Eff_RsPerMkcal","Pct_share"])
    edited_cs = st.experimental_data_editor(cs_df, num_rows="dynamic")

    st.markdown("**Reasons for fuel cost disallowance (breakup)**")
    reasons_df = pd.DataFrame(reasons) if reasons else pd.DataFrame(columns=["reason","amount_rs_cr"])
    edited_reasons = st.experimental_data_editor(reasons_df, num_rows="dynamic")

    st.markdown("---")
    meta_col1, meta_col2, meta_col3 = st.columns(3)
    with meta_col1:
        station_name = st.text_input("Station name", value="KTPS")
    with meta_col2:
        period_label = st.text_input("Report period (e.g. 01-10 Oct 2025)", value="")
    with meta_col3:
        report_date = st.date_input("Report date (pick any date within the 10-day period)", value=datetime.today())

    if st.button("Save / Append this 10-day record"):
        record = {
            "station": station_name,
            "period_label": period_label,
            "report_date": report_date.isoformat(),
            "vc_merc_order": mer_order,
            "vc_recoverable_mod": recoverable,
            "vc_actual": actual_vc,
            "vc_unrecoverable_rpkwh": unrecoverable_rpkwh,
            "fuel_cost_disallowance_rs_cr": disallowance_cr,
            "sp_coal_allowable": sp_allow,
            "sp_coal_actual": sp_actual,
            "aux_power_pct": aux_power,
            "sp_oil_ml_per_kwh": sp_oil,
            "coal_sources": edited_cs.to_json(orient="records"),
            "reasons_breakup": edited_reasons.to_json(orient="records"),
            "raw_text_excerpt": raw_text[:1200]
        }

        # Attempt Google Sheets append first (optional)
        appended_sheet = False
        sheet_msg = ""
        if USE_GSHEETS and GSPREAD_AVAILABLE:
            ok, sheet_msg = gsheet_append_row(record)
            appended_sheet = ok

        # Always append to local CSV fallback
        df_row = pd.DataFrame([record])
        if os.path.exists(DATA_FILE):
            existing = pd.read_csv(DATA_FILE)
            new = pd.concat([existing, df_row], ignore_index=True)
            new.to_csv(DATA_FILE, index=False)
        else:
            df_row.to_csv(DATA_FILE, index=False)

        if appended_sheet:
            st.success("Record appended to Google Sheet AND local CSV.")
        else:
            st.warning("Record appended to local CSV. Google Sheets append not completed: " + str(sheet_msg))

# ---------- Dashboard & Monthly consolidated view ----------
st.header("2) Dashboard / Monthly report")
if os.path.exists(DATA_FILE):
    df = pd.read_csv(DATA_FILE)
    if 'report_date' in df.columns:
        df['report_date'] = pd.to_datetime(df['report_date'])
    st.subheader("Appended 10-day records")
    st.dataframe(df[["station","period_label","report_date","vc_merc_order","vc_recoverable_mod","vc_actual","fuel_cost_disallowance_rs_cr","sp_coal_actual"]].sort_values("report_date"))

    st.subheader("Variable charge comparison (MERC vs Actual)")
    ts = df.sort_values("report_date")
    # ensure numeric columns for plotting
    for c in ["vc_merc_order","vc_actual","vc_recoverable_mod"]:
        if c in ts.columns:
            ts[c] = pd.to_numeric(ts[c], errors="coerce")
    fig = px.line(ts, x="report_date", y=[c for c in ["vc_merc_order","vc_actual","vc_recoverable_mod"] if c in ts.columns], markers=True)
    st.plotly_chart(fig, use_container_width=True)

    st.subheader("Fuel cost disallowance (Rs.Cr) by period")
    fig2 = px.bar(ts, x="period_label", y="fuel_cost_disallowance_rs_cr", hover_data=["report_date"])
    st.plotly_chart(fig2, use_container_width=True)

    st.subheader("Monthly consolidated report (aggregates)")
    df['month'] = df['report_date'].dt.to_period('M')
    agg = df.groupby('month').agg(
        total_disallowance_rs_cr = ("fuel_cost_disallowance_rs_cr","sum"),
        avg_vc_mer = ("vc_merc_order","mean"),
        avg_vc_actual = ("vc_actual","mean"),
        avg_sp_coal_actual = ("sp_coal_actual","mean"),
        avg_aux_power = ("aux_power_pct","mean")
    ).reset_index()
    st.dataframe(agg)

    st.subheader("Top coal sources contributing to high effective cost (combined view)")
    month_sel = st.selectbox("Select month for coal-source rollup", options=sorted(df['month'].astype(str).unique()))
    sub = df[df['month'].astype(str) == month_sel]
    combined = {}
    for cs_json in sub['coal_sources'].dropna():
        try:
            cs_list = pd.read_json(cs_json)
            for _,row in cs_list.iterrows():
                name = str(row.get("name") or "").strip()
                if not name: continue
                eff = row.get("Eff_RsPerMkcal") if "Eff_RsPerMkcal" in row.index else None
                pct = row.get("Pct_share") if "Pct_share" in row.index else None
                if name not in combined:
                    combined[name] = {"count":0, "effs":[], "pct":0.0}
                combined[name]["count"] += 1
                try:
                    if eff: combined[name]["effs"].append(float(eff))
                except: pass
                try:
                    if pct: combined[name]["pct"] += float(pct)
                except: pass
        except:
            pass
    combined_list = []
    for k,v in combined.items():
        avg_eff = sum(v["effs"])/len(v["effs"]) if v["effs"] else None
        combined_list.append({"name":k, "appearances":v["count"], "avg_eff_rs_per_mkcal":avg_eff, "total_pct_share":v["pct"]})
    if combined_list:
        comb_df = pd.DataFrame(combined_list).sort_values(["avg_eff_rs_per_mkcal","appearances"], ascending=[False,False])
        st.dataframe(comb_df)
    else:
        st.info("No coal-source structured data found for selected month.")

    st.subheader("Reasons for fuel-cost disallowance (roll-up)")
    reasons_agg = {}
    for rj in df['reasons_breakup'].dropna():
        try:
            rr = pd.read_json(rj)
            for _,row in rr.iterrows():
                rn = str(row.get("reason") or "").strip()
                amt = row.get("amount_rs_cr")
                try: amt = float(amt) if pd.notna(amt) else 0.0
                except: amt = 0.0
                if rn not in reasons_agg: reasons_agg[rn] = 0.0
                reasons_agg[rn] += amt
        except:
            pass
    if reasons_agg:
        rlst = [{"reason":k,"total_rs_cr":v} for k,v in reasons_agg.items()]
        rdf = pd.DataFrame(rlst).sort_values("total_rs_cr", ascending=False)
        st.dataframe(rdf)
    else:
        st.info("No structured reasons data available yet.")

    if st.button("Export monthly consolidated report (current view) to CSV"):
        out_name = os.path.join(DATA_DIR, f"ktps_monthly_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
        agg.to_csv(out_name, index=False)
        st.success(f"Saved {out_name}")
else:
    st.info("No appended records yet. Upload and save a 10-day report to start building the dashboard.")

st.markdown("""
**Notes**
- Parser is heuristic: review parsed fields before saving.
- Coal-source and reasons tables are stored as JSON strings in the CSV.
- For durable persistence on Streamlit Cloud, configure `USE_GSHEETS=True` and add `gcp_service_account` (full JSON) and `gspread_sheet_id` in Streamlit secrets.
- Requirements: `streamlit`, `pandas`, `python-docx`, `plotly` and optionally `gspread` & `google-auth` if using Sheets.
""")
